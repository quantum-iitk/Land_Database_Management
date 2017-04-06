from tkinter.tix import *
from PIL import ImageTk, Image
import MySQLdb
from base_test import FirstWindow
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENTATION
import pymysql as psql
import os



# install Pillow using pip


db_name = FirstWindow().db_name
root = Tk()
root.title("Land Database Management System")
myContainer = Frame(root)
myContainer.grid()
db = MySQLdb.connect("localhost", port=3306, user="root", db=db_name)
entries = []
variables = []
array = []
new_var = []
numcol = 0


def represents_int(s):
    try:
        int(s)
        return True
    except ValueError:
        return False
# Viewtable

def view_table():
    value = str((listbox1.get(listbox1.curselection())))

    col_window = Toplevel()
    col_window.title("Name of Columns")
    col_window.focus_set()

    cursor1 = db.cursor()
    database1 = db_name
    cursor1.execute(
        "SELECT `COLUMN_NAME` FROM `INFORMATION_SCHEMA`.`COLUMNS` WHERE `TABLE_SCHEMA` = %s AND `TABLE_NAME` = %s",
        (database1, value))
    global numcol, i
    numcol = int(cursor1.rowcount)

    if(numcol>10):
        j=0
        col1=1
        col2=2
        for i in range(0,numcol):

            row1 = [x[0] for x in cursor1.fetchmany()]
            array.insert(i, row1)
            label1 = Label(col_window, text=row1[0])
            label1.grid(row=j, column=col1, sticky=W)
            va = StringVar()
            entry = Entry(col_window, textvariable=StringVar())
            entry.grid(row=j, column=col2, padx=17, pady=17)
            variables.append(va)
            entries.append(entry)
            j=j+1
            if(j==13):
                j=0
                col1=col1+2
                col2=col2+2
            else:
                j=j

    else:
        for i in range(0, numcol):
            row1 = [x[0] for x in cursor1.fetchmany()]

            array.insert(i, row1)
            label1 = Label(col_window, text=row1[0])
            label1.grid(row=i, column=1, sticky=W)
            va = StringVar()
            entry = Entry(col_window, textvariable=StringVar())
            entry.grid(row=i, column=2, padx=2, pady=2)
            variables.append(va)
            entries.append(entry)

    for k in range(0, numcol):
        vary = (", ".join(array[k]))
        new_var.insert(k, vary)


    def insert_data():
        cursor2 = db.cursor()

        # print(entries[j].get())
        p = "("
        q = "("
        for l in range(0, numcol - 1):
            if represents_int(entries[l].get()):
                p = p + entries[l].get() + ","
            else:
                p = p + "'" + entries[l].get() + "'" + ","
            q = q + "`" + new_var[l] + "`" + ","
        if represents_int(entries[numcol - 1].get()):
            p = p + entries[numcol - 1].get() + ")"
        else:
            p = p + "'" + entries[numcol - 1].get() + "'" + ")"
        table_name = "`" + value + "`"

        q = q + "`" + new_var[numcol - 1] + "`" + ")"
        # print (p)
        # print (q)
        cursor2.execute("INSERT INTO {0} {1} VALUES {2} ".format(table_name, q, p))
        db.commit()
        # for x in entries:
        #     print (x.get())

    def delete_data():
        cursor3 = db.cursor()
        table_name = "`" + value + "`"
        r = "`" + new_var[0] + "`"
        if represents_int(entries[0].get()):
            s = entries[0].get()
        else:
            s = "'" + entries[0].get() + "'"

        cursor3.execute("DELETE FROM {0} WHERE {1} ={2} ".format(table_name, r, s))

    def update_data():
        new_var1 = []
        entries1 = []
        cursor4 = db.cursor()
        table_name = "`" + value + "`"

        r = "`" + new_var[0] + "`"

        for m in range(0, numcol):
            if represents_int(entries[m].get()):
                entries1.insert(m, entries[m].get())
            else:
                entries1.insert(m, "'" + entries[m].get() + "'")
        for number in range(0, numcol):
            new_var1.insert(number, "`" + new_var[number] + "`")
        if represents_int(entries[0].get()):
            s = entries[0].get()
        else:
            s = "'" + entries[0].get() + "'"

        for l in range(0, numcol):
            cursor4.execute(
                """UPDATE {0} SET {1}={2} WHERE {3} ={4} """.format(table_name, new_var1[l], entries1[l], r, s))

    def search_data():
        data_container=[]
        new_var1 = []
        cursor5 = db.cursor()
        table_name = "`" + value + "`"

        r = "`" + new_var[0] + "`"
        for number in range(0, numcol):
            new_var1.insert(number, "`" + new_var[number] + "`")

        if represents_int(entries[0].get()):
            s = entries[0].get()
        else:
            s = "'" + entries[0].get() + "'"

        for l in range (0, numcol):
            cursor5.execute("""SELECT {0} FROM {1} WHERE {2} ={3} """.format(new_var1[l],table_name, r, s))
            data_container.insert(l, cursor5.fetchone()[0])

        count = 0
        for data in data_container:
            print(data)
            e = entries[count]
            e.delete(0, END)
            e.insert(count, data)
            count += 1

    if(numcol>10):
        button_insert = Button(col_window, text="INSERT", command=insert_data)
        button_insert.grid(row=15, column=3,)

        button_update = Button(col_window, text="UPDATE", command=update_data)
        button_update.grid(row=15, column=4,)

        button_delete = Button(col_window, text="Delete", command=delete_data)
        button_delete.grid(row=15, column=5)


        button_search = Button(col_window, text="Fetch", command=search_data)
        button_search.grid(row=15, column=6)

    else:
        button_insert = Button(col_window, text="INSERT", command=insert_data)
        button_insert.grid(row=i + 1, column=1, sticky=SW)

        button_update = Button(col_window, text="UPDATE", command=update_data)
        button_update.grid(row=i + 1, column=2, sticky=SW)

        button_delete = Button(col_window, text="Delete", command=delete_data)
        button_delete.grid(row=i + 1, column=3, sticky=SW)


        button_search = Button(col_window, text="Search", command=search_data)
        button_search.grid(row=i+1, column=4)

label = Label(root, text="The tables in the database are as follow:")
label.grid(row=0, column=0, sticky=N)
label.configure(background='green')


# label_print = Label(root, text="Select the listed tables to print:",background='green')
# label_print.grid(row=0, column=3, sticky=E,)

listbox1 = Listbox(root,background='yellow')
listbox1.grid(sticky=NSEW,)


click = Button(root, text="Click to view tables", command=view_table)
click.grid(row=2, column=0, sticky=W, padx=2, pady=2)

my_menu = Menu(root)
menu_bar = Menu(root)
file_menu = Menu(menu_bar, tearoff=0)
# all file menu-items will be added here next
menu_bar.add_cascade(label='Help', menu=file_menu)
root.config(menu=menu_bar)
file_menu = Menu(menu_bar, tearoff=0)
# all file menu-items will be added here next
menu_bar.add_cascade(label='About', menu=file_menu)
root.config(menu=menu_bar)

path = 'E:\\LandDatabase\\WindowsFormsApplication1\\Resources\\Photo.jpg'

# Creates a Tkinter-compatible photo image, which can be used everywhere Tkinter expects an image object.
img = ImageTk.PhotoImage(Image.open(path))

# The Label widget is a standard Tkinter widget used to display a text or image on the screen.
panel = Label(root, image = img)

# The Pack geometry manager packs widgets in rows or columns.
panel.grid(row=1,column=2,sticky=E)
root.configure(background='green')

cursor = db.cursor()
cursor.execute('SELECT table_name FROM information_schema.tables where table_schema=\'' + db_name + '\'')
num_tables = int(cursor.rowcount)

db.commit()


for i in range(0, num_tables):
    row = cursor.fetchone()
    listbox1.insert(END, row[0])
    # lsb_query.insert(END,row[0])

def print_details(*args):
    try:
        value = str((listbox1.get(listbox1.curselection())))
        print_data(value)
    except ValueError:
        pass

printer = Button(root, text="Print", command=print_details)
printer.grid(row=2, column=0, sticky=E, padx=2, pady=2)


def print_data(value):
    connection = psql.connect(host='localhost', user='root', password='', db=db_name, charset='utf8mb4',
                              autocommit=True, cursorclass=psql.cursors.DictCursor)

    cursor = connection.cursor()
    details = db_name.split("_")
    print(details)
    district = details[2]
    tehsil = details[3]
    village = details[4]
    paragana = details[5]

    cursor.execute('SELECT * FROM `%s`'%(value))
    data = cursor.fetchall()

    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    print(section.orientation)

    section.page_width = Inches(18.03)
    section.page_height = Inches(12.76)
    print(section.page_height, section.page_width)
    # /section.orientation = WD_ORIENT.LANDSCAPE
    h = document.add_heading(value.capitalize(), 0)
    h.alignment = 1
    h1=document.add_heading("District: "+district+"   Tehsil:  "+tehsil+"  Village:  "+village+"   Paragana:  "+paragana,0)
    h1.alignment= 1
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    p.alignment =1

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')


    # document.add_picture('monty-truth.png', width=Inches(1.25))

    table = document.add_table(rows=len(data), cols=len(list(data[0].keys())))
    table.style = 'LightShading-Accent1'
    hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    for i,keys in enumerate(list(data[0].keys())):
        hdr_cells[i].text = str(keys)

    for dict in data:
        row_cells = table.add_row().cells
        for i,values in enumerate(list(dict.values())):
            row_cells[i].text = str(values)
    # for item in recordset:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(item.qty)
    #     row_cells[1].text = str(item.id)
    #     row_cells[2].text = item.desc

    document.add_page_break()

    document.save('demo.docx')
    os.startfile('demo.docx','print')

root.mainloop()
db.close()
# Created by Ashutosh Kumar 14148 IIT Kanpur/
