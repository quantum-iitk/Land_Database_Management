from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENTATION
import pymysql as psql
import os


def print_data(form_name,dbname):
    connection = psql.connect(host='localhost', user='root', password='', db=dbname, charset='utf8mb4',
                              autocommit=True, cursorclass=psql.cursors.DictCursor)

    cursor = connection.cursor()
    details = dbname.split("_")
    print(details)
    district = details[2]
    tehsil = details[3]
    village = details[4]
    paragana = details[5]

    cursor.execute('SELECT * FROM `%s`'%(form_name))
    data = cursor.fetchall()

    document = Document()

    section = document.sections[0]
    section.orientation= WD_ORIENTATION.LANDSCAPE
    print(section.orientation)

    section.page_width = Inches(18.03)
    section.page_height = Inches(12.76)
    print(section.page_height,section.page_width)
    # /section.orientation = WD_ORIENT.LANDSCAPE
    h = document.add_heading(form_name.capitalize(), 0)
    h.alignment = 1
    h1=document.add_heading("District: "+district+"   Tehsil:  "+tehsil+"  Village:  "+village+"   Paragana:  "+paragana,0)
    h1.alignment= 1
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    p.alignment =1

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')


    # document.add_picture('monty-truth.png', width=Inches(1.25))

    table = document.add_table(rows=len(data), cols=len(list(data[0].keys())))
    table.style = 'LightShading-Accent1'
    hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    for i,keys in enumerate(list(data[0].keys())):
        hdr_cells[i].text = str(keys)

    for dict in data:
        row_cells = table.add_row().cells
        for i,values in enumerate(list(dict.values())):
            row_cells[i].text = str(values)
    # for item in recordset:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(item.qty)
    #     row_cells[1].text = str(item.id)
    #     row_cells[2].text = item.desc

    document.add_page_break()

    document.save('demo.docx')
    os.startfile('demo.docx','print')